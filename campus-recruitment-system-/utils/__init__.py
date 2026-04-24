# -*- coding: utf-8 -*-
"""
工具类模块 - 包含智能匹配算法、简历解析等核心功能
"""

import os
import re
import json
import jieba
import jieba.analyse
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from docx import Document


# ==================== 文本处理工具 ====================

def chinese_tokenizer(text):
    """
    中文分词器 - 使用jieba进行中文分词
    作为TF-IDF的分词函数
    """
    if not text:
        return []
    # 使用jieba精确模式分词
    words = jieba.lcut(text)
    # 过滤停用词和过短词汇
    stopwords = set(['的', '了', '在', '是', '我', '有', '和', '就', '不', '人',
                     '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去',
                     '你', '会', '着', '没有', '看', '好', '自己', '这', '那'])
    words = [w.strip() for w in words if len(w.strip()) > 1 and w.strip() not in stopwords]
    return words


def extract_keywords(text, topK=20):
    """
    提取文本关键词
    使用jieba的TF-IDF算法提取关键词

    Args:
        text: 输入文本
        topK: 返回前K个关键词

    Returns:
        list: 关键词列表
    """
    if not text:
        return []
    keywords = jieba.analyse.extract_tags(text, topK=topK, withWeight=False)
    return keywords


# ==================== 简历解析工具 ====================

class ResumeParser:
    """
    简历解析类 - 支持PDF和Word文档解析
    提取文本内容用于后续分析
    """

    @staticmethod
    def parse_file(file_path):
        """
        根据文件类型解析简历

        Args:
            file_path: 简历文件路径

        Returns:
            dict: 包含解析结果的字典
        """
        if not os.path.exists(file_path):
            return {'success': False, 'error': '文件不存在'}

        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            if file_ext == '.pdf':
                return ResumeParser._parse_pdf(file_path)
            elif file_ext in ['.doc', '.docx']:
                return ResumeParser._parse_word(file_path)
            else:
                return {'success': False, 'error': f'不支持的文件格式: {file_ext}'}
        except Exception as e:
            return {'success': False, 'error': str(e)}

    @staticmethod
    def _parse_pdf(file_path):
        """
        解析PDF文件
        使用PyPDF2库提取文本
        """
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""

            return {
                'success': True,
                'text': text,
                'pages': len(reader.pages)
            }
        except Exception as e:
            return {'success': False, 'error': f'PDF解析错误: {str(e)}'}

    @staticmethod
    def _parse_word(file_path):
        """
        解析Word文档
        使用python-docx库提取文本
        """
        try:
            doc = Document(file_path)
            text = []

            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text.strip())

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(' '.join(row_text))

            full_text = '\n'.join(text)

            return {
                'success': True,
                'text': full_text,
                'paragraphs': len(doc.paragraphs)
            }
        except Exception as e:
            return {'success': False, 'error': f'Word解析错误: {str(e)}'}

    @staticmethod
    def extract_info(text):
        """
        从简历文本中提取关键信息
        使用正则表达式和关键词匹配

        Args:
            text: 简历文本

        Returns:
            dict: 提取的信息字典
        """
        info = {
            'name': '',
            'phone': '',
            'email': '',
            'education': '',
            'school': '',
            'major': '',
            'skills': []
        }

        if not text:
            return info

        # 提取手机号
        phone_pattern = r'1[3-9]\d{9}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            info['phone'] = phone_match.group()

        # 提取邮箱
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        email_match = re.search(email_pattern, text)
        if email_match:
            info['email'] = email_match.group()

        # 提取学历
        edu_keywords = ['博士', '硕士', '本科', '大专', '研究生', 'MBA', 'EMBA']
        for edu in edu_keywords:
            if edu in text:
                info['education'] = edu
                break

        # 提取技能关键词（常见IT技能）
        skill_keywords = [
            'Python', 'Java', 'C++', 'C#', 'JavaScript', 'Go', 'Rust', 'PHP',
            'HTML', 'CSS', 'React', 'Vue', 'Angular', 'Node.js', 'Django',
            'Flask', 'Spring', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis',
            'Linux', 'Docker', 'Kubernetes', 'AWS', 'Git', '机器学习',
            '深度学习', 'TensorFlow', 'PyTorch', '数据分析', '人工智能',
            'NLP', '计算机视觉', '大数据', 'Hadoop', 'Spark', 'Hive',
            '产品经理', 'UI设计', '测试', '运维', '前端', '后端', '全栈'
        ]

        for skill in skill_keywords:
            if skill.lower() in text.lower() or skill in text:
                info['skills'].append(skill)

        # 去重
        info['skills'] = list(set(info['skills']))

        return info


# ==================== 智能匹配算法核心类 ====================

class ResumeJobMatcher:
    """
    简历与岗位智能匹配类
    基于TF-IDF + 余弦相似度算法实现

    算法原理：
    1. 使用TF-IDF将文本转换为向量表示
    2. 计算两个向量的余弦相似度
    3. 结合多维度规则进行综合评分
    """

    def __init__(self):
        """
        初始化匹配器
        配置jieba词典和TF-IDF向量化器
        """
        # 添加专业词典（可选）
        # jieba.load_userdict('custom_dict.txt')

        # 初始化TF-IDF向量化器
        # tokenizer使用自定义中文分词器
        self.vectorizer = TfidfVectorizer(
            tokenizer=chinese_tokenizer,
            lowercase=False,
            max_features=5000,  # 最大特征数
            min_df=1,  # 最小文档频率
            ngram_range=(1, 2)  # 使用1-gram和2-gram
        )

    def calculate_similarity(self, resume_text, job_text):
        """
        计算简历与岗位的相似度

        Args:
            resume_text: 简历文本
            job_text: 岗位文本

        Returns:
            float: 余弦相似度值（0-1）
        """
        if not resume_text or not job_text:
            return 0.0

        try:
            # 将两篇文本向量化
            tfidf_matrix = self.vectorizer.fit_transform([resume_text, job_text])

            # 计算余弦相似度
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])

            return float(similarity[0][0])
        except Exception as e:
            print(f"相似度计算错误: {e}")
            return 0.0

    def calculate_match_score(self, resume, job):
        """
        综合计算匹配分数
        结合TF-IDF相似度和规则评分

        Args:
            resume: Resume模型对象
            job: Job模型对象

        Returns:
            dict: 包含各维度分数和详细报告
        """
        # 获取完整文本
        resume_text = resume.get_full_text()
        job_text = job.get_full_text()

        # 1. 计算TF-IDF相似度（基础分，权重40%）
        tfidf_similarity = self.calculate_similarity(resume_text, job_text)
        tfidf_score = tfidf_similarity * 40  # 最高40分

        # 2. 学历匹配评分（权重20%）
        education_score = self._calculate_education_score(
            resume.education or '',
            job.education_req or ''
        )

        # 3. 技能匹配评分（权重25%）
        skill_score = self._calculate_skill_score(
            resume.skills or '',
            job.skills_req or ''
        )

        # 4. 专业匹配评分（权重15%）
        major_score = self._calculate_major_score(
            resume.major or '',
            job.major_req or ''
        )

        # 计算总分（0-100）
        final_score = tfidf_score + education_score + skill_score + major_score

        # 生成详细匹配报告
        match_details = self._generate_match_details(
            resume, job,
            tfidf_similarity, education_score, skill_score, major_score
        )

        return {
            'similarity_score': round(tfidf_similarity, 4),  # TF-IDF原始相似度
            'final_score': round(final_score, 2),  # 最终分数（0-100）
            'education_score': round(education_score, 2),  # 学历分（0-20）
            'skill_score': round(skill_score, 2),  # 技能分（0-25）
            'experience_score': round(major_score, 2),  # 专业分（0-15）
            'match_details': json.dumps(match_details, ensure_ascii=False)
        }

    def _calculate_education_score(self, resume_edu, job_edu_req):
        """
        学历匹配评分

        学历等级：博士(5) > 硕士(4) > 本科(3) > 大专(2) > 其他(1)
        如果简历学历 >= 岗位要求，得满分；否则按比例扣分
        """
        edu_levels = {
            '博士': 5, '硕士研究生': 4, '硕士': 4, '研究生': 4,
            '本科': 3, '大学本科': 3, '学士': 3,
            '大专': 2, '专科': 2, '高职': 2,
            '高中': 1, '中专': 1, '其他': 0
        }

        max_score = 20  # 学历满分20分

        if not job_edu_req:
            return max_score  # 无要求则给满分

        # 提取学历等级
        resume_level = 0
        job_level = 0

        for edu, level in edu_levels.items():
            if edu in resume_edu:
                resume_level = max(resume_level, level)
            if edu in job_edu_req:
                job_level = max(job_level, level)

        # 如果简历学历 >= 岗位要求，得满分
        if resume_level >= job_level and job_level > 0:
            return max_score
        elif resume_level > 0 and job_level > 0:
            # 按比例给分
            return max_score * (resume_level / job_level)

        return max_score * 0.5  # 无法判断时给一半分

    def _calculate_skill_score(self, resume_skills, job_skills_req):
        """
        技能匹配评分

        计算简历技能与岗位要求的重合度
        每个匹配技能加分，最高25分
        """
        max_score = 25

        if not job_skills_req:
            return max_score * 0.8  # 无要求给80%

        if not resume_skills:
            return 0

        # 提取技能关键词
        resume_skill_list = [s.strip().lower() for s in re.split(r'[,，;；、\s]+', resume_skills) if s.strip()]
        job_skill_list = [s.strip().lower() for s in re.split(r'[,，;；、\s]+', job_skills_req) if s.strip()]

        if not job_skill_list:
            return max_score * 0.8

        # 计算匹配的技能数
        matched = 0
        for job_skill in job_skill_list:
            for resume_skill in resume_skill_list:
                if job_skill in resume_skill or resume_skill in job_skill:
                    matched += 1
                    break

        # 匹配率 = 匹配数 / 要求总数
        match_rate = matched / len(job_skill_list) if job_skill_list else 0

        # 基础分 + 匹配奖励分
        score = max_score * 0.3 + max_score * 0.7 * match_rate

        return min(score, max_score)

    def _calculate_major_score(self, resume_major, job_major_req):
        """
        专业匹配评分

        检查简历专业是否符合岗位要求
        """
        max_score = 15

        if not job_major_req:
            return max_score * 0.8

        if not resume_major:
            return 0

        # 提取专业关键词
        job_majors = [m.strip() for m in re.split(r'[,，;；、\s]+', job_major_req) if m.strip()]

        if not job_majors:
            return max_score * 0.8

        # 检查是否匹配
        resume_major_clean = resume_major.strip().lower()
        for major in job_majors:
            if major.lower() in resume_major_clean or resume_major_clean in major.lower():
                return max_score

        # 相关专业匹配（简化处理）
        related_keywords = ['计算机', '软件', '信息', '电子', '通信', '自动化']
        for kw in related_keywords:
            if kw in resume_major and any(kw in m for m in job_majors):
                return max_score * 0.8

        return max_score * 0.3

    def _generate_match_details(self, resume, job, tfidf_sim, edu_score, skill_score, major_score):
        """
        生成详细的匹配报告

        包含：
        - 匹配项列表
        - 不匹配项列表
        - 优势分析
        - 劣势分析
        - 改进建议
        """
        details = {
            'matched_items': [],
            'unmatched_items': [],
            'advantages': [],
            'disadvantages': [],
            'suggestions': []
        }

        # 学历匹配详情
        if edu_score >= 18:
            details['matched_items'].append(f"学历达标：{resume.education}")
            details['advantages'].append("学历符合或超过岗位要求")
        elif edu_score >= 10:
            details['matched_items'].append(f"学历接近：{resume.education}")
            details['suggestions'].append("如有更高学历可补充说明")
        else:
            details['unmatched_items'].append(f"学历不匹配：当前{resume.education}，要求{job.education_req}")
            details['disadvantages'].append("学历未达到岗位要求")
            details['suggestions'].append("考虑学历提升或投递学历要求较低的岗位")

        # 技能匹配详情
        resume_skills = [s.strip() for s in re.split(r'[,，;；、\s]+', resume.skills or '') if s.strip()]
        job_skills = [s.strip() for s in re.split(r'[,，;；、\s]+', job.skills_req or '') if s.strip()]

        matched_skills = []
        missing_skills = []

        for job_skill in job_skills:
            is_matched = False
            for resume_skill in resume_skills:
                if job_skill.lower() in resume_skill.lower() or resume_skill.lower() in job_skill.lower():
                    matched_skills.append(job_skill)
                    is_matched = True
                    break
            if not is_matched:
                missing_skills.append(job_skill)

        if matched_skills:
            details['matched_items'].append(f"匹配技能：{', '.join(matched_skills[:5])}")
            details['advantages'].append(f"掌握{len(matched_skills)}项岗位所需技能")

        if missing_skills:
            details['unmatched_items'].append(f"缺失技能：{', '.join(missing_skills[:5])}")
            if len(missing_skills) > len(matched_skills):
                details['disadvantages'].append("技能匹配度较低")
            details['suggestions'].append(f"建议学习：{', '.join(missing_skills[:3])}")

        # 专业匹配详情
        if major_score >= 12:
            details['matched_items'].append(f"专业匹配：{resume.major}")
        elif resume.major:
            details['unmatched_items'].append(f"专业差异：{resume.major} vs {job.major_req}")

        # 综合建议
        final_score = tfidf_sim * 40 + edu_score + skill_score + major_score
        if final_score >= 80:
            details['suggestions'].insert(0, "匹配度很高，建议尽快投递")
        elif final_score >= 60:
            details['suggestions'].insert(0, "匹配度良好，可以尝试投递")
        elif final_score >= 40:
            details['suggestions'].insert(0, "匹配度一般，针对性优化简历后再投递")
        else:
            details['suggestions'].insert(0, "匹配度较低，建议寻找更匹配的岗位")

        return details


# ==================== 初始化匹配器实例 ====================

matcher = ResumeJobMatcher()


def get_matcher():
    """
    获取全局匹配器实例
    """
    return matcher
